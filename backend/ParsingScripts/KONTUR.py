import os
import time
import shutil
import re
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from io import BytesIO 
import requests
from docx import Document
import asyncio
from concurrent.futures import ThreadPoolExecutor

import pandas as pd

BASE_URL = "https://www.kontur-extern.ru/price-download/77"
DOWNLOAD_DIR = os.environ.get("DOWNLOAD_DIR")
SELENIUM_URL = os.environ.get("SELENIUM_URL")
DOCTRANSLATOR_URL = os.environ.get("DOCTRANSLATOR_URL")

_regions = [
    ("01", "Республика Адыгея"),
    ("02", "Республика Башкортостан"),
    ("03", "Республика Бурятия"),
    ("04", "Республика Алтай"),
]

regions = [
    ("01", "Республика Адыгея"),
    ("02", "Республика Башкортостан"),
    ("03", "Республика Бурятия"),
    ("04", "Республика Алтай"),
    ("05", "Республика Дагестан"),
    ("06", "Республика Ингушетия"),
    ("07", "Республика Кабардино-Балкария"),
    ("08", "Республика Калмыкия"),
    ("09", "Республика Карачаево-Черкесия"),
    ("10", "Республика Карелия"),
    ("11", "Республика Коми"),
    ("12", "Республика Марий Эл"),
    ("13", "Республика Мордовия"),
    ("14", "Республика Саха (Якутия)"),
    ("15", "Республика Северная Осетия - Алания"),
    ("16", "Республика Татарстан"),
    ("17", "Республика Тыва"),
    ("18", "Республика Удмуртия"),
    ("19", "Республика Хакасия"),
    ("20", "Республика Чечня"),
    ("21", "Республика Чувашия"),
    ("22", "Алтайский край"),
    ("23", "Краснодарский край"),
    ("24", "Красноярский край"),
    ("25", "Приморский край"),
    ("26", "Ставропольский край"),
    ("27", "Хабаровский край"),
    ("28", "Амурская область"),
    ("29", "Архангельская область"),
    ("30", "Астраханская область"),
    ("31", "Белгородская область"),
    ("32", "Брянская область"),
    ("33", "Владимирская область"),
    ("34", "Волгоградская область"),
    ("35", "Вологодская область"),
    ("36", "Воронежская область"),
    ("37", "Ивановская область"),
    ("38", "Иркутская область"),
    ("39", "Калининградская область"),
    ("40", "Калужская область"),
    ("41", "Камчатский край"),
    ("42", "Кемеровская область"),
    ("43", "Кировская область"),
    ("44", "Костромская область"),
    ("45", "Курганская область"),
    ("46", "Курская область"),
    ("47", "Ленинградская область"),
    ("48", "Липецкая область"),
    ("49", "Магаданская область"),
    ("50", "Московская область"),
    ("51", "Мурманская область"),
    ("52", "Нижегородская область"),
    ("53", "Новгородская область"),
    ("54", "Новосибирская область"),
    ("55", "Омская область"),
    ("56", "Оренбургская область"),
    ("57", "Орловская область"),
    ("58", "Пензенская область"),
    ("59", "Пермский край"),
    ("60", "Псковская область"),
    ("61", "Ростовская область"),
    ("62", "Рязанская область"),
    ("63", "Самарская область"),
    ("64", "Саратовская область"),
    ("65", "Сахалинская область"),
    ("66", "Свердловская область"),
    ("67", "Смоленская область"),
    ("68", "Тамбовская область"),
    ("69", "Тверская область"),
    ("70", "Томская область"),
    ("71", "Тульская область"),
    ("72", "Тюменская область"),
    ("73", "Ульяновская область"),
    ("74", "Челябинская область"),
    ("75", "Забайкальский край"),
    ("76", "Ярославская область"),
    ("77", "Москва"),
    ("78", "Санкт-Петербург"),
    ("79", "Еврейская АО"),
    ("83", "Ненецкий АО"),
    ("86", "Ханты-Мансийский АО"),
    ("87", "Чукотский АО"),
    ("89", "Ямало-Ненецкий АО"),
    ("90", "Запорожская область"),
    ("91", "Республика Крым"),
    ("92", "Севастополь"),
    ("93", "Донецкая народная республика"),
    ("94", "Луганская народная республика"),
    ("95", "Херсонская область"),
    ("99", "Байконур")
]

def download_price_file(driver, wait):
    def download_selenium():
        try:
            link = wait.until(EC.element_to_be_clickable((
                By.XPATH, "//a[contains(text(), 'Скачать полный прайс-лист, часть 2')]"
            )))
            driver.execute_script("arguments[0].click();", link)
            return True
        except Exception as e:
            print(f"    ❌ Ошибка при скачивании файла: {e}")
        return False

    def download_requests():
        try:
            print("загрузка из requests")
            link = wait.until(EC.element_to_be_clickable((
                By.XPATH, "//a[contains(text(), 'Скачать полный прайс-лист, часть 2')]"
            )))

            file_url = link.get_attribute("href")
            response = requests.get(file_url)
            
            if response.status_code == 200:
                print("Файл успешно загружен")
                path = os.path.join(DOWNLOAD_DIR, f"pricelist.doc")
                with open(path, "wb") as f:
                    f.write(response.content)
                return True
            else:
                print("Ошибка при загрузке файла")
        except Exception as e:
            print(f"    ❌ Ошибка при скачивании файла: {e}")
        return False
    
    download_selenium()
    start_time = time.time()
    download_finished = False
    while not download_finished:
        time.sleep(2)
        files_list = os.listdir(DOWNLOAD_DIR)
        print(files_list)
        if len(files_list)>0:
            download_finished = True
        if time.time() - start_time > 30:
            print("❌ Время ожидания истекло")
            break
            
    if not download_finished:
        download_finished = download_requests()

    return download_finished   

def parse_docx_from_bytes(docx_bytes: BytesIO):
    try:
        docx_bytes.seek(0)
        doc = Document(docx_bytes)

        ip_usn = ip_osno = ul_usn = ul_osno = "❌"
        budget_plus = budget = "❌"

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                joined = " ".join(row_text).lower()

                if "оптимальный плюс" in joined and "1 год" in joined:
                    ip_usn = row_text[2] if len(row_text) > 2 else "❌"
                    ip_osno = row_text[3] if len(row_text) > 3 else "❌"
                    ul_usn = row_text[4] if len(row_text) > 4 else "❌"
                    ul_osno = row_text[5] if len(row_text) > 5 else "❌"

                if "бюджетник плюс" in joined and "1 год" in joined:
                    for item in reversed(row_text):
                        if any(ch.isdigit() for ch in item):
                            budget_plus = item
                            break

                if "бюджетник" in joined and "1 год" in joined and "плюс" not in joined and "максимальный" not in joined:
                    for item in reversed(row_text):
                        if any(ch.isdigit() for ch in item):
                            budget = item
                            break

        return ip_usn, ul_usn, ip_osno, ul_osno, budget_plus, budget

    except Exception as e:
        return [f"Ошибка: {e}"] * 6

def extract_common_prices_from_bytes(docx_bytes:BytesIO):
    try:
        docx_bytes.seek(0)
        doc = Document(docx_bytes)

        target_keys = {
            "1+4": "1+4",
            "1+9": "1+9",
            "1+19": "1+19",
            "1+49": "1+49",
            "1+99": "1+99",
            "1+199": "1+199",
            "1+499": "1+499"
        }

        common_prices = {key: "❌" for key in target_keys}
        found_table = False

        for table in doc.tables:
            if not table.rows or not table.rows[0].cells:
                continue

            first_cell = table.rows[0].cells[0].text.strip()
            if "Количество организаций" not in first_cell:
                continue

            found_table = True
            print("\n📊 Найдена таблица с количеством организаций. Анализ строк:")

            for row in table.rows[1:]:  # пропускаем заголовок
                try:
                    col1 = row.cells[0].text.strip()
                    col2 = row.cells[1].text.strip()

                    print(f"  🔹 Строка: [{col1}] -> [{col2}]")

                    for key in target_keys:
                        if key in col1 and common_prices[key] == "❌":
                            cleaned = re.sub(r"[^\d]", "", col2)
                            if cleaned:
                                common_prices[key] = int(cleaned)
                                print(f"    ✅ Найдено: {key} = {common_prices[key]}")
                            else:
                                print(f"    ⚠ Не удалось распознать цену для {key}: [{col2}]")
                except Exception as e:
                    print(f"  ⚠ Ошибка при обработке строки: {e}")
                    continue

            break

        if not found_table:
            print("❗ Таблица с 'Количество организаций на обслуживании' не найдена.")

        return [
            common_prices["1+4"],
            common_prices["1+9"],
            common_prices["1+19"],
            common_prices["1+49"],
            common_prices["1+99"],
            common_prices["1+199"],
            common_prices["1+499"],
        ]

    except Exception as e:
        return [f"Ошибка: {e}"] * 7

def translate_file(file_path,file_name):
    with open(file_path, "rb") as f:
        url = DOCTRANSLATOR_URL
        files = {"file": (file_name, f, "multipart/form-data")}
        response = requests.post(url, files=files)

    response.raise_for_status()

    file = BytesIO(response.content)

    return file

def init_webdriver(local = False, headless = True):
    
    options = webdriver.ChromeOptions()
    if headless:
        options.add_argument('--headless=new')
        options.add_argument('--disable-gpu')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36")
    options.add_argument("--lang=ru-RU")
    options.add_argument("--window-size=1920,1080")
    options.add_argument("--disable-extensions")
    options.add_argument("--disable-plugins")
    options.add_argument("--disable-popup-blocking")
    options.add_argument('--start-maximized')

    prefs = {
        "download.default_directory": DOWNLOAD_DIR,
        "download.prompt_for_download": False,
        "download.directory_upgrade": True,
        "safebrowsing.enabled": True
    }
    options.add_experimental_option("prefs", prefs)


    if local:
        driver = webdriver.Chrome(options=options)
    else:
        driver = webdriver.Remote(
            command_executor=SELENIUM_URL,
            options=options
        )

    if headless:
        driver.execute_cdp_cmd(
            "Browser.setDownloadBehavior",
            {
                "behavior": "allow",
                "downloadPath": DOWNLOAD_DIR
            }
        )

        
    return driver


def init_dataframe():
    df = pd.DataFrame(columns=["Код региона", "Название региона", "ИП (УСН)", "ИП (ОСНО)", "ЮЛ (УСН)", "ЮЛ (ОСНО)", "Бюджетник плюс", "Бюджетник",
                               "1+4", "1+9", "1+19", "1+49", "1+99", "1+199", "1+499"])
    return df


def read_data_row(region_id,region_name):
    latest_file = max([os.path.join(DOWNLOAD_DIR, f) for f in os.listdir(DOWNLOAD_DIR)], key=os.path.getctime)
    target_path = os.path.join(DOWNLOAD_DIR, f"{region_id}_pricelist.doc")
    if os.path.abspath(latest_file) != os.path.abspath(target_path):
        shutil.move(latest_file, target_path)

    print(f"✔ Скачан: {target_path}")

    file_docx = translate_file(target_path,f"{region_id}_pricelist.doc")

    # Извлекаем основные тарифы
    ip_usn, ul_usn, ip_osno, ul_osno, budget_plus, budget = parse_docx_from_bytes(file_docx)
    print(f"  └ ИП (УСН): {ip_usn}")
    print(f"  └ ЮЛ (УСН): {ul_usn}")
    print(f"  └ ИП (ОСНО): {ip_osno}")
    print(f"  └ ЮЛ (ОСНО): {ul_osno}")
    print(f"  └ Бюджетник плюс: {budget_plus}")
    print(f"  └ Бюджетник: {budget}")

    # Извлекаем тарифы "Общий"
    common_prices = extract_common_prices_from_bytes(file_docx)
    print(f"  └ 1+4: {common_prices[0]}")
    print(f"  └ 1+9: {common_prices[1]}")
    print(f"  └ 1+19: {common_prices[2]}")
    print(f"  └ 1+49: {common_prices[3]}")
    print(f"  └ 1+99: {common_prices[4]}")
    print(f"  └ 1+199: {common_prices[5]}")
    print(f"  └ 1+499: {common_prices[6]}")

    # Добавляем данные в Excel
    row = [int(region_id), region_name, ip_usn, ip_osno, ul_usn, ul_osno, budget_plus, budget] + common_prices
    return row
def main():
    total_regions = len(regions)

    print("test")
    driver = init_webdriver(local=False,headless=True)

    wait = WebDriverWait(driver, 20) 

    df = init_dataframe()

    # === Основной цикл ===
    for idx, (region_id, region_name) in enumerate(regions, 1):
        print(f"\n➡ Обрабатываем: {region_id} – {region_name} ({idx}/{total_regions})")

        try:
            driver.get(BASE_URL.replace("77", region_id))

   

            for f in os.listdir(DOWNLOAD_DIR):
                try:
                    os.remove(os.path.join(DOWNLOAD_DIR, f))
                except:
                    pass
            wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))

            if download_price_file(driver, wait):
                df.loc[len(df)] = read_data_row(region_id,region_name)
                # print(df)
                # time.sleep(10)
            else: print(f"  ❌ Не удалось скачать файл с тарифами для")
            



        except Exception as e:
            print(f"❌ Ошибка в регионе {region_id} – {region_name}: {e}")

        print(f"✅ Завершено для {region_name} ({round(idx / total_regions * 100, 1)}%)")

    # === Завершение ===
    driver.quit()

    return df


async def async_selenium():
    loop = asyncio.get_running_loop() 
    with ThreadPoolExecutor() as executor:
        df = await loop.run_in_executor(executor, main)
    return df


async def parse_kontur():
    df = await async_selenium()
    return df

if __name__ == "__main__":
    for key, value in dict(os.environ).items():
        print(f"{key}: {value}")
    df = main()
    df.to_excel("text.xlsx")